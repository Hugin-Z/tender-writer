# -*- coding: utf-8 -*-
from __future__ import annotations

import json
import os
import shutil
import subprocess
import sys
import tempfile
import unittest
from pathlib import Path

from docx import Document


REPO_ROOT = Path(__file__).resolve().parents[1]


class TenderWriterRegressionTest(unittest.TestCase):
    maxDiff = None

    def make_workspace(self) -> Path:
        temp_dir = tempfile.mkdtemp(prefix="tender_writer_test_")
        self.addCleanup(lambda: shutil.rmtree(temp_dir, ignore_errors=True))
        workspace = Path(temp_dir) / "tender-writer"
        shutil.copytree(
            REPO_ROOT,
            workspace,
            ignore=shutil.ignore_patterns(".venv", "__pycache__", "*.pyc", "output"),
        )
        (workspace / "assets" / ".ingest_history.json").write_text("{}", encoding="utf-8")
        return workspace

    def run_script(self, workspace: Path, *args: str) -> subprocess.CompletedProcess[str]:
        env = dict(os.environ)
        env["PYTHONUTF8"] = "1"
        result = subprocess.run(
            [sys.executable, "-X", "utf8", *args],
            cwd=workspace,
            capture_output=True,
            text=True,
            encoding="utf-8",
            errors="replace",
            env=env,
        )
        if result.returncode != 0:
            self.fail(
                "command failed\n"
                f"cwd: {workspace}\n"
                f"cmd: {' '.join(args)}\n"
                f"stdout:\n{result.stdout}\n"
                f"stderr:\n{result.stderr}"
            )
        return result

    def test_scoring_outline_and_compliance_workflow(self) -> None:
        workspace = self.make_workspace()
        output_dir = workspace / "output"
        output_dir.mkdir(exist_ok=True)

        brief_text = """# 招标文件摘要

## 评分办法

### 商务评分
- 类似业绩（4分）：每提供 1 个相关案例得 2 分，最高 4 分。
- 项目经理资质（2分）：具备高级职称得 2 分。

### 技术评分
- 技术方案（6分）：方案完整、针对性强、实施路径清晰得 6 分。
- 售后服务（2分）：响应时限明确、培训计划完整得 2 分。
"""
        brief_path = output_dir / "tender_brief.md"
        brief_path.write_text(brief_text, encoding="utf-8")

        self.run_script(workspace, "scripts/build_scoring_matrix.py", str(brief_path))
        matrix_path = output_dir / "scoring_matrix.csv"
        self.assertTrue(matrix_path.exists())
        matrix_text = matrix_path.read_text(encoding="utf-8-sig")
        self.assertIn("类似业绩", matrix_text)
        self.assertIn("技术方案", matrix_text)

        self.run_script(workspace, "scripts/generate_outline.py", str(matrix_path))
        outline_path = output_dir / "outline.md"
        self.assertTrue(outline_path.exists())
        outline_text = outline_path.read_text(encoding="utf-8")
        self.assertIn("技术方案", outline_text)
        self.assertIn("覆盖度自检", outline_text)

        docx_path = output_dir / "tender_response.docx"
        self.run_script(
            workspace,
            "scripts/docx_builder.py",
            "--out",
            str(docx_path),
            "--project",
            "测试项目",
            "--bidder",
            "测试投标人",
        )
        self.assertTrue(docx_path.exists())

        chapter_path = output_dir / "chapter_01.md"
        chapter_path.write_text(
            """# 第一章 响应摘要

## 商务响应
- 类似业绩：已提供 2 个相关案例作为证明材料。
- 项目经理资质：项目经理具备高级职称，并附证书。

## 技术与服务响应
- 技术方案：方案完整、针对性强，包含实施路径和交付计划。
- 售后服务：承诺 30 分钟响应、4 小时到场，并提供培训计划。
""",
            encoding="utf-8",
        )
        self.run_script(workspace, "scripts/append_chapter.py", str(docx_path), str(chapter_path))

        self.run_script(workspace, "scripts/compliance_check.py", str(docx_path), str(matrix_path))
        report_path = output_dir / "compliance_report.md"
        self.assertTrue(report_path.exists())
        report_text = report_path.read_text(encoding="utf-8")
        self.assertIn("覆盖度检查", report_text)
        self.assertIn("格式检查", report_text)

    def test_asset_ingest_and_triage_workflow(self) -> None:
        workspace = self.make_workspace()

        self.run_script(
            workspace,
            "scripts/add_company.py",
            "Apply Partner Ltd",
            "partner",
            "--alias",
            "TestPartner",
        )

        performance_doc = workspace / "assets" / "类似业绩" / "own_default" / "_inbox" / "sample_performance.docx"
        performance_doc.parent.mkdir(parents=True, exist_ok=True)
        doc = Document()
        doc.add_paragraph("Smart City Platform Project")
        doc.add_paragraph("采购人：测试市大数据局")
        doc.add_paragraph("合同金额：128万元")
        doc.add_paragraph("项目时间：2025年03月12日")
        doc.add_paragraph("项目内容包括数据治理、驾驶舱、事件协同处置。")
        doc.save(performance_doc)

        ingest_result = self.run_script(workspace, "scripts/ingest_assets.py", "业绩", "own_default")
        self.assertIn("处理成功: 1", ingest_result.stdout)
        self.assertIn("Smart_City_Platform_Project.md", ingest_result.stdout)

        history = json.loads((workspace / "assets" / ".ingest_history.json").read_text(encoding="utf-8"))
        self.assertTrue(
            any(
                row.get("company_id") == "own_default"
                and row.get("category") == "业绩"
                and row.get("target", "").endswith("Smart_City_Platform_Project.md")
                for row in history.values()
            )
        )

        unsorted_dir = workspace / "_inbox_unsorted"
        unsorted_dir.mkdir(exist_ok=True)

        resume_doc = unsorted_dir / "sample_resume.docx"
        doc = Document()
        doc.add_paragraph("张三")
        doc.add_paragraph("项目经理简历")
        doc.add_paragraph("学历：本科")
        doc.add_paragraph("职称：高级工程师")
        doc.add_paragraph("工作经历：10年政务信息化项目经验")
        doc.save(resume_doc)

        cert_doc = unsorted_dir / "TestPartner_cert.docx"
        doc = Document()
        doc.add_paragraph("TestPartner 质量管理体系认证证书")
        doc.add_paragraph("证书编号：QA-2026-002")
        doc.add_paragraph("发证机关：认证机构")
        doc.add_paragraph("有效期至：2027年06月30日")
        doc.save(cert_doc)

        triage_preview = self.run_script(workspace, "scripts/triage_unsorted.py")
        self.assertIn("类别建议: 简历", triage_preview.stdout)
        self.assertIn("类别建议: 资质", triage_preview.stdout)
        self.assertIn("partner_testpartner", triage_preview.stdout)

        triage_apply = self.run_script(workspace, "scripts/triage_unsorted.py", "--apply")
        self.assertIn("跳过（未识别公司归属）", triage_apply.stdout)
        self.assertIn("已触发摄入", triage_apply.stdout)

        history = json.loads((workspace / "assets" / ".ingest_history.json").read_text(encoding="utf-8"))
        self.assertTrue(
            any(
                row.get("company_id") == "partner_testpartner"
                and row.get("category") == "资质"
                and row.get("target", "").endswith(".md")
                for row in history.values()
            )
        )


if __name__ == "__main__":
    unittest.main(verbosity=2)
