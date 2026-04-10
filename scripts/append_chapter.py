# -*- coding: utf-8 -*-
"""
append_chapter.py · 将 markdown 章节追加到现有 docx（阶段 4 辅助脚本）

输入:
    output/tender_response.docx
    章节 markdown 文件

能力:
    - 识别 # / ## / ### / #### 标题
    - 识别普通段落
    - 识别 - / * 无序列表

限制:
    - 暂不解析 markdown 表格、图片、复杂嵌套结构
"""

from __future__ import annotations

import argparse
import sys
from pathlib import Path

try:
    from docx import Document
except ImportError:
    print("[错误] 缺少 python-docx 依赖，请先双击 install.bat 安装依赖。", file=sys.stderr)
    sys.exit(1)

from docx_builder import add_chapter, add_paragraph, apply_default_styles, set_page_margins


def flush_paragraph(document: Document, buffer: list[str]) -> None:
    if not buffer:
        return
    text = " ".join(part.strip() for part in buffer if part.strip()).strip()
    if text:
        add_paragraph(document, text)
    buffer.clear()


def append_markdown(document: Document, markdown_text: str) -> None:
    paragraph_buffer: list[str] = []

    for raw_line in markdown_text.splitlines():
        line = raw_line.rstrip()
        stripped = line.strip()

        if not stripped:
            flush_paragraph(document, paragraph_buffer)
            continue

        if stripped.startswith("#"):
            flush_paragraph(document, paragraph_buffer)
            level = len(stripped) - len(stripped.lstrip("#"))
            level = min(max(level, 1), 4)
            title = stripped[level:].strip()
            add_chapter(document, title, level=level)
            continue

        if stripped.startswith("- ") or stripped.startswith("* "):
            flush_paragraph(document, paragraph_buffer)
            bullet_text = stripped[2:].strip()
            add_paragraph(document, f"- {bullet_text}", first_line_indent_chars=0)
            continue

        paragraph_buffer.append(stripped)

    flush_paragraph(document, paragraph_buffer)


def main() -> None:
    parser = argparse.ArgumentParser(
        description="将 markdown 章节追加到现有 docx",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog=__doc__,
    )
    parser.add_argument("docx_path", help="现有 docx 路径")
    parser.add_argument("chapter_md", help="章节 markdown 文件路径")
    args = parser.parse_args()

    docx_path = Path(args.docx_path)
    chapter_md = Path(args.chapter_md)

    if not docx_path.exists():
        print(f"[错误] 找不到 docx 文件: {docx_path}", file=sys.stderr)
        sys.exit(1)
    if not chapter_md.exists():
        print(f"[错误] 找不到 markdown 文件: {chapter_md}", file=sys.stderr)
        sys.exit(1)

    document = Document(str(docx_path))
    apply_default_styles(document)
    set_page_margins(document)
    append_markdown(document, chapter_md.read_text(encoding="utf-8"))
    document.save(str(docx_path))

    print(f"[完成] 已将章节追加到: {docx_path}")


if __name__ == "__main__":
    main()
