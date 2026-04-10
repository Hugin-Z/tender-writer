# -*- coding: utf-8 -*-
"""
compliance_check.py · 合规终审脚本（阶段 5）

输入:
    output/tender_response.docx
    output/scoring_matrix.csv

输出:
    output/compliance_report.md

说明:
    这是自动化初筛工具，目标是尽量把明显的漏答、模板残留和格式风险提出来。
    它不能替代标书经理或法务的人工终审。
"""

from __future__ import annotations

import argparse
import csv
import re
import sys
from pathlib import Path

TEMPLATE_RESIDUES = [
    "xxx 公司",
    "XXX公司",
    "XX公司",
    "xxx 项目",
    "XXX项目",
    "XX项目",
    # 占位符模式必须带尖括号或方括号,避免误伤正文里合法出现的"甲方单位"等词
    "<甲方单位>",
    "<甲方名称>",
    "<乙方名称>",
    "<项目名称>",
    "<投标人名称>",
    "[甲方单位]",
    "[甲方名称]",
    "[乙方名称]",
    "[项目名称]",
    "[投标人名称]",
    "【甲方单位】",
    "【项目名称】",
    "【投标人名称】",
    "【公司名称】",
    "TODO",
    "todo",
    "【待补充】",
    "【待填写】",
    "【待人工确认】",
    "示例文本",
]

SUBSTANTIAL_KEYWORDS = ["完全响应", "完全满足", "实质性响应", "无偏离", "满足要求"]
FORMAT_CHECKS = {
    "封面": ["投标文件", "项目名称", "投标人"],
    "目录": ["目录"],
}
NORMALIZE_PATTERN = re.compile(r"[\s\u3000，,。；;：:（）()\[\]【】\-—_/]")


def normalize_for_match(text: str) -> str:
    return NORMALIZE_PATTERN.sub("", text)


def read_docx_text(docx_path: Path) -> tuple[str, dict]:
    try:
        from docx import Document
    except ImportError:
        print("[错误] 缺少 python-docx 依赖，请先双击 install.bat 安装依赖。", file=sys.stderr)
        sys.exit(1)

    document = Document(str(docx_path))
    paragraph_texts: list[str] = []
    paragraph_styles: list[str] = []

    for para in document.paragraphs:
        text = para.text.strip()
        if text:
            paragraph_texts.append(text)
            paragraph_styles.append(getattr(para.style, "name", ""))

    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                paragraph_texts.append(" | ".join(cells))

    full_text = "\n".join(paragraph_texts)
    format_info = {
        "paragraph_count": len(document.paragraphs),
        "non_empty_paragraphs": len(paragraph_texts),
        "heading_count": sum(1 for name in paragraph_styles if name.startswith("Heading")),
        "cover_preview": paragraph_texts[:12],
        "section_margins": [],
    }

    for section in document.sections:
        format_info["section_margins"].append({
            "top_cm": round(section.top_margin.cm, 2) if section.top_margin else None,
            "bottom_cm": round(section.bottom_margin.cm, 2) if section.bottom_margin else None,
            "left_cm": round(section.left_margin.cm, 2) if section.left_margin else None,
            "right_cm": round(section.right_margin.cm, 2) if section.right_margin else None,
        })

    return full_text, format_info


def load_matrix(matrix_csv: Path) -> list[dict]:
    rows: list[dict] = []
    with open(matrix_csv, "r", encoding="utf-8-sig", newline="") as handle:
        reader = csv.DictReader(handle)
        for row in reader:
            rows.append(row)
    return rows


def extract_candidates(row: dict) -> list[str]:
    item = (row.get("评分项") or "").strip()
    keywords_field = (row.get("关键词") or "").strip()

    candidates: list[str] = []
    if item:
        candidates.append(item)
    if keywords_field:
        candidates.extend([part.strip() for part in re.split(r"[/、，,；;\s]+", keywords_field) if part.strip()])

    for token in re.split(r"[（）()\[\]【】、，,；;：:\s]+", item):
        token = token.strip()
        if 2 <= len(token) <= 10:
            candidates.append(token)

    unique: list[str] = []
    seen: set[str] = set()
    for candidate in candidates:
        normalized = normalize_for_match(candidate)
        if len(normalized) < 2 or normalized in seen:
            continue
        seen.add(normalized)
        unique.append(candidate)
    return unique


def check_coverage(docx_text: str, matrix: list[dict]) -> list[dict]:
    normalized_doc = normalize_for_match(docx_text)
    results: list[dict] = []

    for row in matrix:
        item = (row.get("评分项") or "").strip()
        score = (row.get("分值") or "").strip()
        chapter = (row.get("应答章节") or "").strip()
        candidates = extract_candidates(row)

        exact_hit = item and normalize_for_match(item) in normalized_doc
        matched_candidates = [
            candidate for candidate in candidates
            if normalize_for_match(candidate) in normalized_doc
        ]

        if exact_hit or len(matched_candidates) >= 2:
            status = "covered"
        elif len(matched_candidates) == 1:
            status = "partial"
        else:
            status = "missing"

        results.append({
            "item": item,
            "score": score,
            "chapter": chapter,
            "status": status,
            "candidates": candidates,
            "matched_candidates": matched_candidates,
        })

    return results


def check_substantial_response(docx_text: str, matrix: list[dict]) -> list[dict]:
    results: list[dict] = []
    lowered_doc = docx_text

    for row in matrix:
        item = (row.get("评分项") or "").strip()
        risk = (row.get("风险提示") or "").strip()
        if not any(marker in item or marker in risk for marker in ("★", "▲", "实质性响应")):
            continue

        candidates = extract_candidates(row)
        responded = False
        evidence = ""

        for candidate in candidates:
            start = 0
            while True:
                idx = lowered_doc.find(candidate, start)
                if idx == -1:
                    break
                window = lowered_doc[max(0, idx - 80): idx + len(candidate) + 120]
                matched_phrase = next((phrase for phrase in SUBSTANTIAL_KEYWORDS if phrase in window), "")
                if matched_phrase:
                    responded = True
                    evidence = f"{candidate} 附近出现“{matched_phrase}”"
                    break
                start = idx + len(candidate)
            if responded:
                break

        if not responded:
            global_phrase = next((phrase for phrase in SUBSTANTIAL_KEYWORDS if phrase in lowered_doc), "")
            if global_phrase:
                evidence = f"全文出现过“{global_phrase}”，但未能与该条款建立近邻关联"
            else:
                evidence = "未找到明确响应表述"

        results.append({
            "item": item,
            "responded": responded,
            "evidence": evidence,
        })

    return results


def check_template_residues(docx_text: str) -> list[str]:
    found: list[str] = []
    for residue in TEMPLATE_RESIDUES:
        if residue in docx_text:
            idx = docx_text.find(residue)
            context = docx_text[max(0, idx - 20): idx + len(residue) + 20].replace("\n", " ")
            found.append(f"残留“{residue}” 上下文: ...{context}...")
    return found


def check_format(docx_text: str, format_info: dict) -> list[str]:
    issues: list[str] = []
    expected_margins = {"top_cm": 2.54, "bottom_cm": 2.54, "left_cm": 3.17, "right_cm": 3.17}

    margins = format_info.get("section_margins", [])
    if not margins:
        issues.append("未能读取到 section 页边距信息")
    else:
        for index, section_margin in enumerate(margins):
            for key, expected in expected_margins.items():
                actual = section_margin.get(key)
                if actual is None:
                    issues.append(f"section[{index}] 缺少 {key}")
                elif abs(actual - expected) > 0.05:
                    issues.append(f"section[{index}] {key} 实际 {actual} cm，期望 {expected} cm")

    for check_name, keywords in FORMAT_CHECKS.items():
        if not all(keyword in docx_text for keyword in keywords):
            issues.append(f"缺少明显的“{check_name}”特征关键词: {', '.join(keywords)}")

    if format_info.get("heading_count", 0) == 0:
        issues.append("未检测到 Heading 样式标题，目录和层级结构可能未按 docx_builder 规范生成")

    return issues


def render_report(
    coverage: list[dict],
    substantial: list[dict],
    residues: list[str],
    format_issues: list[str],
    docx_path: Path,
    matrix_csv: Path,
) -> str:
    total = len(coverage)
    covered = [row for row in coverage if row["status"] == "covered"]
    partial = [row for row in coverage if row["status"] == "partial"]
    missing = [row for row in coverage if row["status"] == "missing"]

    sub_total = len(substantial)
    sub_responded = [row for row in substantial if row["responded"]]
    sub_missing = [row for row in substantial if not row["responded"]]

    lines: list[str] = []
    lines.append("# 合规终审报告（compliance_report）")
    lines.append("")
    lines.append(f"> **被检 docx**: {docx_path}")
    lines.append(f"> **评分矩阵**: {matrix_csv}")
    lines.append("> **说明**: 本报告是自动化初筛结果，不能替代标书经理和法务的人工终审。")
    lines.append("")
    lines.append("---")
    lines.append("")
    lines.append("## 一、评分项覆盖度检查")
    lines.append("")
    lines.append(f"- 评分项总数: {total}")
    lines.append(f"- 明确覆盖: {len(covered)}")
    lines.append(f"- 弱覆盖: {len(partial)}")
    lines.append(f"- 未覆盖: {len(missing)}")
    lines.append(f"- 覆盖率: {(len(covered) / total * 100) if total else 0:.1f}%")
    lines.append("")

    lines.append("### 未覆盖清单（必须返回阶段 4 补写）")
    lines.append("")
    if missing:
        for row in missing:
            preview = "，".join(row["candidates"][:5])
            lines.append(f"- ❌ **{row['item']}**（分值 {row['score']}）")
            lines.append(f"  候选关键词: {preview}")
            lines.append(f"  目标章节: {row['chapter']}")
    else:
        lines.append("- ✅ 未发现明确漏答项")
    lines.append("")

    lines.append("### 弱覆盖清单（建议人工复核）")
    lines.append("")
    if partial:
        for row in partial:
            matched = "，".join(row["matched_candidates"])
            lines.append(f"- ⚠️ **{row['item']}** 仅命中关键词: {matched}")
    else:
        lines.append("- ✅ 未发现弱覆盖项")
    lines.append("")

    lines.append("---")
    lines.append("")
    lines.append("## 二、★/▲ 条款响应检查")
    lines.append("")
    lines.append(f"- ★/▲ 条款总数: {sub_total}")
    lines.append(f"- 明确响应: {len(sub_responded)}")
    lines.append(f"- 未明确响应: {len(sub_missing)}")
    lines.append("")

    lines.append("### 未明确响应清单")
    lines.append("")
    if sub_missing:
        for row in sub_missing:
            lines.append(f"- ❌ **{row['item']}**：{row['evidence']}")
    else:
        if sub_total == 0:
            lines.append("- ⚠️ 评分矩阵中未标出任何 ★/▲ 条款，请人工复核 tender_brief.md")
        else:
            lines.append("- ✅ 所有 ★/▲ 条款均检测到近邻响应表述")
    lines.append("")

    lines.append("---")
    lines.append("")
    lines.append("## 三、模板残留检查")
    lines.append("")
    if residues:
        for residue in residues:
            lines.append(f"- ❌ {residue}")
    else:
        lines.append("- ✅ 未检测到常见模板残留（如 XXX 公司、TODO、【待补充】）")
    lines.append("")

    lines.append("---")
    lines.append("")
    lines.append("## 四、格式检查")
    lines.append("")
    if format_issues:
        for issue in format_issues:
            lines.append(f"- ⚠️ {issue}")
    else:
        lines.append("- ✅ 页边距、封面/目录特征、标题样式未发现明显问题")
    lines.append("")

    serious_issues = len(missing) + len(sub_missing) + len(residues)
    lines.append("---")
    lines.append("")
    lines.append("## 五、结论")
    lines.append("")
    if serious_issues == 0:
        lines.append("- ✅ 自动检查未发现严重问题，但仍需完成人工终审后再提交。")
    else:
        lines.append(
            f"- ❌ 自动检查发现 {serious_issues} 项严重问题"
            f"（漏答 {len(missing)} 项、★/▲ 未明确响应 {len(sub_missing)} 项、模板残留 {len(residues)} 项），"
            "必须返回阶段 4 修正后重新检查。"
        )

    return "\n".join(lines) + "\n"


def main() -> None:
    parser = argparse.ArgumentParser(
        description="合规终审（阶段 5）",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog=__doc__,
    )
    parser.add_argument("docx_path", help="最终标书 docx 路径")
    parser.add_argument("matrix_csv", help="评分矩阵 CSV 路径")
    parser.add_argument("--out", default="output", help="报告输出目录，默认 output/")
    args = parser.parse_args()

    docx_path = Path(args.docx_path)
    matrix_csv = Path(args.matrix_csv)

    if not docx_path.exists():
        print(f"[错误] 找不到 docx 文件: {docx_path}", file=sys.stderr)
        sys.exit(1)
    if not matrix_csv.exists():
        print(f"[错误] 找不到评分矩阵: {matrix_csv}", file=sys.stderr)
        sys.exit(1)

    out_dir = Path(args.out)
    out_dir.mkdir(parents=True, exist_ok=True)

    print(f"[信息] 正在读取 docx: {docx_path}")
    docx_text, format_info = read_docx_text(docx_path)
    print(f"[信息] 正在读取评分矩阵: {matrix_csv}")
    matrix = load_matrix(matrix_csv)

    coverage = check_coverage(docx_text, matrix)
    substantial = check_substantial_response(docx_text, matrix)
    residues = check_template_residues(docx_text)
    format_issues = check_format(docx_text, format_info)

    report_path = out_dir / "compliance_report.md"
    report_path.write_text(
        render_report(coverage, substantial, residues, format_issues, docx_path, matrix_csv),
        encoding="utf-8",
    )

    print(f"[完成] 合规报告已写入: {report_path}")
    print()
    print("=" * 60)
    print(f"自动检查统计: 明确漏答 {sum(1 for row in coverage if row['status'] == 'missing')} 项，"
          f"弱覆盖 {sum(1 for row in coverage if row['status'] == 'partial')} 项，"
          f"★/▲ 未明确响应 {sum(1 for row in substantial if not row['responded'])} 项，"
          f"模板残留 {len(residues)} 项，格式问题 {len(format_issues)} 项")
    print("=" * 60)


if __name__ == "__main__":
    main()
