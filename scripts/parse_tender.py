# -*- coding: utf-8 -*-
"""
parse_tender.py · 招标文件解析脚本(阶段 1)

用途:
    把招标文件(PDF 或 docx)解析成结构化的 tender_brief.json 和
    tender_brief.md。识别招标文件中关键章节(评分办法、资格要求、
    实质性响应、废标条款、格式要求等)并定位。

使用方法:
    通过项目根目录的 run_script.bat 调用,不要直接调用 python:
        run_script.bat parse_tender.py "<招标文件绝对路径>"
        run_script.bat parse_tender.py "<招标文件绝对路径>" --out output

参数:
    tender_path : 招标文件路径,支持 .pdf 和 .docx
    --out       : 输出目录,默认为当前工作目录下的 output/

输出:
    <out>/tender_brief.json   结构化 JSON 解析结果
    <out>/tender_brief.md     基于 templates/tender_brief.md 填充的 markdown
"""

import argparse
import json
import os
import re
import sys
from pathlib import Path

# 中文章节关键词到内部 key 的映射
SECTION_KEYWORDS = {
    "project_info": ["项目概况", "项目简介", "项目背景", "项目说明", "采购概况"],
    "budget": ["预算金额", "最高限价", "采购预算", "项目预算", "投标控制价"],
    "duration": ["工期", "服务期", "实施周期", "项目周期", "交付时间"],
    "qualification": ["资格要求", "投标人资格", "供应商资格", "资质要求"],
    "scoring": ["评分办法", "评标办法", "评审办法", "评分标准", "评审标准"],
    "substantial_response": ["实质性响应", "实质性要求", "实质性条款", "★条款", "▲条款"],
    "disqualification": ["废标条款", "废标条件", "无效投标", "拒绝投标"],
    "format_requirement": ["格式要求", "投标文件格式", "装订要求", "编制要求"],
    "submission": ["投标文件构成", "投标文件组成", "应答文件", "投标文件递交"],
    "open_bid": ["开标时间", "开标地点", "投标截止时间", "递交投标文件"],
}

# 资质要求常见关键词,用于在资格要求章节内进一步识别具体项
QUAL_PATTERNS = {
    "营业执照": r"营业执照",
    "注册资金": r"注册资[本金][^\n]{0,30}",
    "ISO9001": r"ISO\s*9001",
    "ISO27001": r"ISO\s*27001",
    "ISO20000": r"ISO\s*20000",
    "CMMI": r"CMMI[\s\-]?[2345]?",
    "信息系统集成资质": r"信息系统集成[\s\S]{0,10}资[质格]",
    "等保": r"等[级保]保护[\s\S]{0,10}[二三四]级",
    "类似业绩": r"类似[\s\S]{0,5}业绩",
    "项目经理资质": r"项目经理[\s\S]{0,30}",
}


def read_pdf(path: Path) -> str:
    """读取 PDF 文件,返回纯文本(每页之间用 \n\n 分隔)"""
    try:
        import pdfplumber
    except ImportError:
        print("[错误] 缺少 pdfplumber 依赖。请先双击 install.bat 安装依赖。", file=sys.stderr)
        sys.exit(1)

    pages_text = []
    with pdfplumber.open(str(path)) as pdf:
        for i, page in enumerate(pdf.pages):
            text = page.extract_text() or ""
            pages_text.append(text)
    return "\n\n".join(pages_text)


def read_docx(path: Path) -> str:
    """读取 docx 文件,返回纯文本(段落用 \n 分隔,表格用制表符分隔)"""
    try:
        from docx import Document
    except ImportError:
        print("[错误] 缺少 python-docx 依赖。请先双击 install.bat 安装依赖。", file=sys.stderr)
        sys.exit(1)

    doc = Document(str(path))
    parts = []
    # 段落
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text)
    # 表格
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            parts.append("\t".join(cells))
    return "\n".join(parts)


def normalize_text(text: str) -> str:
    """文本预处理:繁简转换兜底、去除多余空白"""
    try:
        from opencc import OpenCC
        cc = OpenCC("t2s")
        text = cc.convert(text)
    except Exception:
        # 没装 opencc 或转换失败时,跳过繁简转换
        pass
    # 把多个连续空白压成单空格,但保留换行
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text


def locate_sections(text: str) -> dict:
    """
    在文本中按 SECTION_KEYWORDS 定位每个关键章节,
    返回 {section_key: {"start": int, "end": int, "title": str, "content": str}}
    """
    lines = text.split("\n")
    # 找出每个关键词出现的起始行号
    hits = []  # [(line_no, section_key, matched_keyword)]
    for idx, line in enumerate(lines):
        line_stripped = line.strip()
        for sec_key, kws in SECTION_KEYWORDS.items():
            for kw in kws:
                # 章节标题通常较短(<= 30 字)且包含关键词
                if kw in line_stripped and len(line_stripped) <= 50:
                    hits.append((idx, sec_key, kw, line_stripped))
                    break

    # 去重:同一行只保留第一个匹配
    seen_lines = set()
    unique_hits = []
    for h in hits:
        if h[0] in seen_lines:
            continue
        seen_lines.add(h[0])
        unique_hits.append(h)

    # 按行号排序,并把每个 hit 的 end 设为下一个 hit 的 start
    unique_hits.sort(key=lambda x: x[0])
    sections = {}
    for i, (line_no, sec_key, kw, title) in enumerate(unique_hits):
        end_line = unique_hits[i + 1][0] if i + 1 < len(unique_hits) else len(lines)
        content = "\n".join(lines[line_no:end_line]).strip()
        # 同一 section_key 可能多次命中,只保留最长的内容(通常是正章节而非目录条目)
        if sec_key not in sections or len(content) > len(sections[sec_key]["content"]):
            sections[sec_key] = {
                "start": line_no,
                "end": end_line,
                "title": title,
                "matched_keyword": kw,
                "content": content,
            }
    return sections


def extract_budget(text: str) -> str:
    """从全文中尝试提取预算金额(返回原文片段)"""
    patterns = [
        r"(预算金额|最高限价|采购预算|投标控制价)[::\s]*([0-9.,]+\s*万?元?)",
        r"([0-9.,]+\s*万元)",
    ]
    for p in patterns:
        m = re.search(p, text)
        if m:
            return m.group(0)
    return ""


def extract_duration(text: str) -> str:
    """从全文中尝试提取工期"""
    patterns = [
        r"(工期|服务期|实施周期)[::\s]*([0-9]+\s*[个]?\s*[日月年天])",
        r"([0-9]+\s*个月)",
    ]
    for p in patterns:
        m = re.search(p, text)
        if m:
            return m.group(0)
    return ""


def extract_qualifications(qual_section_text: str) -> list:
    """从资格要求章节中识别具体的资质项"""
    found = []
    if not qual_section_text:
        return found
    for name, pat in QUAL_PATTERNS.items():
        m = re.search(pat, qual_section_text)
        if m:
            found.append({"name": name, "snippet": m.group(0)})
    return found


def extract_substantial_marks(text: str) -> list:
    """提取所有 ★ / ▲ 标记的条款(取所在行作为上下文)"""
    marks = []
    for idx, line in enumerate(text.split("\n")):
        if "★" in line or "▲" in line:
            marks.append({"line_no": idx, "text": line.strip()})
    return marks


def parse_tender(tender_path: Path) -> dict:
    """主解析函数,返回结构化的 dict"""
    suffix = tender_path.suffix.lower()
    if suffix == ".pdf":
        raw = read_pdf(tender_path)
    elif suffix == ".docx":
        raw = read_docx(tender_path)
    else:
        print(f"[错误] 不支持的文件格式:{suffix}。仅支持 .pdf 和 .docx", file=sys.stderr)
        sys.exit(1)

    text = normalize_text(raw)
    sections = locate_sections(text)
    qual_text = sections.get("qualification", {}).get("content", "")

    result = {
        "source_file": str(tender_path),
        "char_count": len(text),
        "sections": {k: {"title": v["title"], "matched_keyword": v["matched_keyword"],
                         "content": v["content"], "content_preview": v["content"][:500]}
                     for k, v in sections.items()},
        "extracted": {
            "budget": extract_budget(text),
            "duration": extract_duration(text),
            "qualifications": extract_qualifications(qual_text),
            "substantial_response_marks": extract_substantial_marks(text)[:50],
        },
        "raw_text": text,
    }
    return result


def render_brief_md(result: dict, template_path: Path) -> str:
    """
    基于 templates/tender_brief.md 模板渲染填充结果。
    模板中未能自动填充的字段,保留 "【待补充】" 占位,提示用户人工核查。
    """
    sections = result.get("sections", {})
    extracted = result.get("extracted", {})

    def get_section_preview(key):
        sec = sections.get(key)
        if not sec:
            return "【未在招标文件中定位到该章节,请人工补充】"
        full_text = sec.get("content", "").strip()
        if full_text:
            return full_text[:500]
        return sec.get("content_preview", "").strip() or "【章节内容为空,请人工核查】"

    qual_lines = []
    quals = extracted.get("qualifications", [])
    if quals:
        for q in quals:
            qual_lines.append(f"- **{q['name']}**:{q['snippet']}")
    else:
        qual_lines.append("- 【未在资格要求章节中识别到常见资质项,请人工补充】")

    sub_lines = []
    marks = extracted.get("substantial_response_marks", [])
    if marks:
        for m in marks[:30]:
            sub_lines.append(f"- (行 {m['line_no']}) {m['text']}")
        if len(marks) > 30:
            sub_lines.append(f"- ……(共 {len(marks)} 条 ★/▲ 标记,此处仅显示前 30 条)")
    else:
        sub_lines.append("- 【未检测到 ★/▲ 标记的实质性响应条款,请人工核查招标文件原文】")

    md = f"""# 招标文件解读简报(tender_brief)

> **来源文件**:{result.get('source_file', '')}
> **字符总数**:{result.get('char_count', 0)}
> **生成时间**:由 parse_tender.py 自动生成

> ⚠️ 本文档是后续所有阶段的**唯一事实来源(single source of truth)**。
> 凡是标注为"【待补充】"或"【未在招标文件中定位到】"的字段,
> **必须由用户人工核查并补充**,严禁让模型凭记忆脑补。

---

## 一、项目基本信息

- **项目名称**:【待补充,请从招标文件首页/封面提取】
- **采购人**:【待补充】
- **采购代理机构**:【待补充】
- **项目编号**:【待补充】
- **预算金额**:{extracted.get('budget', '') or '【未自动识别,请人工补充】'}
- **总工期**:{extracted.get('duration', '') or '【未自动识别,请人工补充】'}
- **关键里程碑**:【待补充】

---

## 二、投标人资格要求

{chr(10).join(qual_lines)}

> 章节原文摘要(前 500 字符):
>
> {get_section_preview('qualification')}

---

## 三、评分办法

> ⚠️ 这是技术标编制最关键的章节,请人工通读章节原文并填写下表。

### 评分维度权重

- 技术分:【__】 分
- 商务分:【__】 分
- 价格分:【__】 分
- 合计:100 分

### 章节原文摘要(前 500 字符)

{get_section_preview('scoring')}

---

## 四、实质性响应条款(★/▲)

{chr(10).join(sub_lines)}

> 章节原文摘要(前 500 字符):
>
> {get_section_preview('substantial_response')}

---

## 五、废标条款

{get_section_preview('disqualification')}

---

## 六、格式要求

{get_section_preview('format_requirement')}

> 关键格式参数(请人工核查并补充):
> - 字体字号:【待补充】
> - 行距:【待补充】
> - 页边距:【待补充】
> - 封面样式:【待补充】
> - 目录要求:【待补充】

---

## 七、投标文件构成与装订

{get_section_preview('submission')}

> - 正本份数:【待补充】
> - 副本份数:【待补充】
> - 电子版份数:【待补充】
> - 装订方式:【待补充】

---

## 八、关键时间节点

{get_section_preview('open_bid')}

> - 投标截止时间:【待补充】
> - 开标时间:【待补充】
> - 开标地点:【待补充】

---

## 九、人工核查 checklist(请用户逐项打钩)

- [ ] 项目基本信息已补全(名称、预算、工期、里程碑)
- [ ] 资格要求已逐项核对,我方满足
- [ ] 评分办法各维度权重已填写
- [ ] 所有 ★/▲ 条款已列出并标注响应方式
- [ ] 废标条款已通读,关键风险点已标注
- [ ] 格式要求已记录(字号、行距、页边距)
- [ ] 投标文件构成与份数已确认
- [ ] 投标截止时间已记录,留出充足缓冲

---

> ✅ 本简报经用户确认无误后,才能进入阶段 2(评分矩阵构建)。
"""
    return md


def main():
    parser = argparse.ArgumentParser(
        description="招标文件解析(阶段 1)",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog=__doc__,
    )
    parser.add_argument("tender_path", help="招标文件路径(.pdf 或 .docx)")
    parser.add_argument("--out", default="output", help="输出目录,默认 output/")
    args = parser.parse_args()

    tender_path = Path(args.tender_path)
    if not tender_path.exists():
        print(f"[错误] 找不到招标文件:{tender_path}", file=sys.stderr)
        sys.exit(1)

    out_dir = Path(args.out)
    out_dir.mkdir(parents=True, exist_ok=True)

    print(f"[信息] 正在解析:{tender_path}")
    result = parse_tender(tender_path)

    # 写 JSON(去掉 raw_text 以减小体积,raw_text 单独写一个文件供调试)
    json_result = {k: v for k, v in result.items() if k != "raw_text"}
    json_path = out_dir / "tender_brief.json"
    with open(json_path, "w", encoding="utf-8") as f:
        json.dump(json_result, f, ensure_ascii=False, indent=2)
    print(f"[完成] JSON 已写入:{json_path}")

    raw_path = out_dir / "tender_raw.txt"
    with open(raw_path, "w", encoding="utf-8") as f:
        f.write(result["raw_text"])
    print(f"[完成] 原文文本已写入:{raw_path}")

    # 写 markdown
    skill_root = Path(__file__).resolve().parent.parent
    template_path = skill_root / "templates" / "tender_brief.md"
    md_text = render_brief_md(result, template_path)
    md_path = out_dir / "tender_brief.md"
    with open(md_path, "w", encoding="utf-8") as f:
        f.write(md_text)
    print(f"[完成] markdown 已写入:{md_path}")

    print()
    print("=" * 60)
    print("阶段 1 完成。下一步:")
    print(f"  1. 用户人工 review {md_path}")
    print("  2. 补全所有【待补充】字段")
    print("  3. 用户确认后,进入阶段 2:")
    print("     run_script.bat build_scoring_matrix.py output/tender_brief.md")
    print("=" * 60)


if __name__ == "__main__":
    main()
